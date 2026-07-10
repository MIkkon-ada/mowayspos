import io
import json
import re
from collections import Counter, defaultdict
from datetime import date, timedelta

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy import and_, or_, text
from sqlalchemy.orm import Session

from .. import crud, models
from ..database import get_db
from ..domain import submission_status as SS
from ..permissions import (
    PROJECT_ROLE_COLLABORATOR,
    PROJECT_ROLE_COORDINATOR,
    PROJECT_ROLE_OWNER,
    can_access_confirmation_center,
    can_view_issue_decisions,
    can_view_issue_risks,
    get_all_project_roles,
    get_current_user_name,
    get_user_context_from_db,
    require_login,
    require_project_access,
)
from ..services.project_resolution import resolve_project_context

router = APIRouter(prefix="/api/dashboard", tags=["dashboard"])

# ── 任务状态分组 ─────────────────────────────────────────────
_NOT_STARTED  = {"未开始"}
_IN_PROGRESS  = {"推进中", "进行中"}
_COMPLETED    = {"已完成"}
_DELAYED      = {"延期"}
_PAUSED       = {"暂缓", "搁置"}

# ── 提交状态分组：统一从 domain.submission_status 导入 ────────
# 不再在此处定义散 set，确保 dashboard 与AI 确认中心数字口径一致。
_SUB_PENDING   = SS.PENDING_OWNER_REVIEW
_SUB_RETURNED  = SS.RETURNED_TO_SUBMITTER
_SUB_CONFIRMED = SS.CONFIRMED_AND_STORED
_CEO_PENDING   = SS.WAITING_CEO_DECISION
# CEO已批示 = CEO 已给指示，等待 owner 确认入库（属于 PENDING_OWNER_REVIEW）
# _CEO_DECIDED 不再独立计数，避免与 _SUB_PENDING 双重计数。
# 如需展示"CEO已批示"计数，使用 SS.stats_ceo_decided()。


# ── plan_time 时间工具 ───────────────────────────────────────

def _month_to_iso(month: str | None) -> str | None:
    """将前端传来的 '2026年6月' 转为 plan_time LIKE 用的 '2026-06' 前缀；已是 ISO 格式则原样返回。"""
    if not month:
        return month
    m = re.match(r'(\d{4})年(\d{1,2})月', month.strip())
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}"
    return month


def _parse_plan_ym(plan_time: str):
    """
    解析 plan_time 字符串，返回 (start_ym, end_ym)，单位是 year*12+month。
    支持："2026-06"、"2026年6月"、"2026-05~2026-07"、"2026年5月~2026年8月"。
    解析失败返回 None。
    """
    if not plan_time:
        return None
    # 区间格式
    m = re.search(r'(\d{4})[年\-](\d{1,2})[月]?[^0-9]+(\d{4})[年\-](\d{1,2})', plan_time)
    if m:
        return (int(m[1]) * 12 + int(m[2]), int(m[3]) * 12 + int(m[4]))
    # 单月格式
    m = re.search(r'(\d{4})[年\-](\d{1,2})', plan_time)
    if m:
        ym = int(m[1]) * 12 + int(m[2])
        return (ym, ym)
    return None


def _is_overdue(plan_time: str) -> bool:
    """计划截止月已过且任务未完成时返回 True（用于自动识别超期）。"""
    r = _parse_plan_ym(plan_time)
    if not r:
        return False
    today = date.today()
    return r[1] < today.year * 12 + today.month


def _plan_time_in_current_month(plan_time: str) -> bool:
    """当前月在 plan_time 范围内时返回 True。"""
    r = _parse_plan_ym(plan_time)
    if not r:
        return False
    today = date.today()
    cur_ym = today.year * 12 + today.month
    return r[0] <= cur_ym <= r[1]


# ── 内部工具 ─────────────────────────────────────────────────

def _get_visible_project_ids(context: dict, db: Session) -> list[int] | None:
    """
    返回当前用户可见的 project_id 列表。
    管理员返回 None（无限制）；普通用户返回合并了新旧两套权限系统的 ID 列表。
    """
    if context["can_view_all"]:
        return None
    ids: set[int] = set()
    # 新系统：project_members 表
    person_id = context.get("person_id")
    if person_id:
        rows = db.execute(
            text("SELECT DISTINCT project_id FROM project_members WHERE person_id = :pid"),
            {"pid": person_id},
        ).fetchall()
        ids.update(r[0] for r in rows)
    # 旧系统：Person 字符串字段推导的项目名 → 转成 ID
    visible_names = context.get("visible_projects") or []
    if visible_names:
        rows = db.query(models.Project.id).filter(models.Project.name.in_(visible_names)).all()
        ids.update(r[0] for r in rows)
    return list(ids)


def _apply_project_scope(query, context, model_cls, visible_project_ids=None):
    """全局模式：按权限限制可见数据，同时兼容新(project_id)和旧(special_project)字段。"""
    if context["can_view_all"]:
        return query
    visible_names = context.get("visible_projects") or []
    ids = list(visible_project_ids or [])
    conds = []
    sp_col = getattr(model_cls, "special_project", None)
    if visible_names and sp_col is not None:
        conds.append(sp_col.in_(visible_names))
    if ids and hasattr(model_cls, "project_id"):
        conds.append(model_cls.project_id.in_(ids))
    if not conds:
        return query.filter(False)
    return query.filter(or_(*conds))


def _proj_or_filter(model_cls, project_id: int, proj_name: str | None, str_col: str = "special_project"):
    """
    新旧数据兼容过滤器：
    - 新数据：project_id 列 == project_id
    - 旧数据：project_id IS NULL AND <str_col> == proj_name
    """
    col = getattr(model_cls, str_col, None)
    if proj_name and col is not None:
        return or_(
            model_cls.project_id == project_id,
            and_(model_cls.project_id.is_(None), col == proj_name),
        )
    return model_cls.project_id == project_id


def _effective_roles(context: dict, project_id: int, proj_name: str | None, db) -> set[str]:
    """
    计算当前用户在指定项目的有效角色集合。
    super_admin 返回全量集合表示无限制。
    project_members 未迁移时回落 ctx["project_roles"] 旧字符串逻辑。
    """
    if context.get("is_tech_admin"):
        return {"super_admin", "owner", "coordinator", "member", "project_ceo"}

    person_id = context.get("person_id")
    roles: set[str] = set()

    if person_id:
        db_roles = get_all_project_roles(person_id, project_id, db)
        roles.update(db_roles)

    # project_members 未迁移，回落旧字符串逻辑
    if not roles and proj_name:
        old_role = context.get("project_roles", {}).get(proj_name)
        if old_role == PROJECT_ROLE_OWNER:
            roles.add("owner")
        elif old_role == PROJECT_ROLE_COORDINATOR:
            roles.add("coordinator")
        elif old_role == PROJECT_ROLE_COLLABORATOR:
            roles.add("member")

    # project_ceo 只来自真实 project_members/project_roles，不再由全局 is_ceo 回落。
    if proj_name and proj_name in context.get("ceo_projects", []):
        roles.add("project_ceo")

    return roles


def _task_stats(tasks: list) -> dict:
    statuses = [t.status or "" for t in tasks]
    # 自动超期：计划截止月已过 + 状态仍是进行中/未开始（避免与手动标记"延期"重复计数）
    _active = _NOT_STARTED | _IN_PROGRESS
    auto_overdue = sum(
        1 for t in tasks
        if (t.status or "") in _active and _is_overdue(t.plan_time or "")
    )
    return {
        "total_tasks":  len(tasks),
        "not_started":  sum(1 for s in statuses if s in _NOT_STARTED),
        "in_progress":  sum(1 for s in statuses if s in _IN_PROGRESS),
        "completed":    sum(1 for s in statuses if s in _COMPLETED),
        "delayed":      sum(1 for s in statuses if s in _DELAYED) + auto_overdue,
        "paused":       sum(1 for s in statuses if s in _PAUSED),
    }


def _submission_stats(subs: list) -> dict:
    return {
        "total_submissions":          len(subs),
        "pending_owner_confirmation": SS.stats_pending_owner(subs),
        "returned_submissions":       SS.stats_returned(subs),
        "confirmed_submissions":      SS.stats_confirmed(subs),
    }


def _ceo_decision_stats(subs: list) -> dict:
    return {
        "pending_ceo_decisions": SS.stats_waiting_ceo(subs),
        # CEO已批示：CEO 已给指示、owner 尚未入库；属于 pending_owner 视图的子集，
        # 此处单独展示供 CEO 角色查看"我已批示了多少条"，不与 pending_owner_confirmation 重复计入总量。
        "ceo_decided_awaiting_owner": SS.stats_ceo_decided(subs),
    }


def _empty_project_overview(context: dict, label: str) -> dict:
    """special_project 无法解析时返回空统计，不报 500。"""
    return {
        "project": {"id": None, "name": label},
        "access": {
            "can_view_decisions": False,
            "can_view_risks": False,
            "can_view_confirmation_center": False,
            "can_view_settings": bool(context.get("can_view_settings")),
            "effective_roles": [],
        },
        "task_stats":       {"total_tasks": 0, "not_started": 0, "in_progress": 0, "completed": 0, "delayed": 0, "paused": 0},
        "achievement_stats": {"total_achievements": 0, "recent_achievements": []},
        "issue_stats":      {"total_issues": 0, "open_issues": 0, "high_priority_issues": 0, "waiting_ceo_decision": 0},
        "submission_stats": {"total_submissions": 0, "pending_owner_confirmation": 0, "returned_submissions": 0, "confirmed_submissions": 0},
        "ceo_decision_stats": {"pending_ceo_decisions": 0, "ceo_decided_awaiting_owner": 0},
        "recent":           {"submissions": [], "tasks": [], "issues": []},
        # Legacy fields
        "summary":          {"task_count": 0, "achievement_count": 0, "open_issue_count": 0, "pending_confirmation_count": 0},
        "project_cards":    [],
        "status_stats":     {},
        "pending_tasks":    [],
        "decisions":        [],
        "risks":            [],
        "latest_achievements": [],
    }


# ── 项目模式 overview ─────────────────────────────────────────

def _require_global_read_scope(context: dict) -> None:
    if not (context.get("is_tech_admin") or context.get("is_ceo")):
        raise HTTPException(403, "permission denied")


def _project_overview(
    context: dict,
    project_id: int,
    proj_name: str | None,
    owner_filter: str | None,
    status_filter: str | None,
    month_filter: str | None,
    db: Session,
) -> dict:
    roles     = _effective_roles(context, project_id, proj_name, db)
    is_super  = "super_admin" in roles
    is_owner  = "owner" in roles
    is_project_ceo = "project_ceo" in roles
    is_global_ceo = bool(context.get("is_ceo"))
    is_coord  = "coordinator" in roles
    # member only: has member role but none of the higher roles
    is_member = "member" in roles and not is_owner and not is_coord and not is_project_ceo and not is_super

    # ── Tasks ──────────────────────────────────────────────
    task_q = db.query(models.Task).filter(
        _proj_or_filter(models.Task, project_id, proj_name),
        models.Task.is_deleted == False,
    )
    if is_member:
        # member 只看自己负责的任务
        task_q = task_q.filter(models.Task.owner == context["name"])
    if owner_filter:
        task_q = task_q.filter(models.Task.owner == owner_filter)
    if status_filter:
        task_q = task_q.filter(models.Task.status == status_filter)
    if month_filter:
        task_q = task_q.filter(models.Task.plan_time.like(f"{_month_to_iso(month_filter)}%"))
    tasks = task_q.order_by(models.Task.updated_at.desc()).all()

    # ── Issues ─────────────────────────────────────────────
    issue_q = db.query(models.Issue).filter(
        _proj_or_filter(models.Issue, project_id, proj_name)
    )
    issues = issue_q.order_by(models.Issue.updated_at.desc()).all()

    open_issues   = [i for i in issues if (i.status or "") in ("待处理", "处理中")]
    high_pri      = [i for i in open_issues if (i.priority or "") == "高"]
    ceo_issues    = [i for i in issues if bool(i.need_decision_by) or "决策" in (i.issue_type or "")]

    # ── Achievements ───────────────────────────────────────
    ach_q = db.query(models.Achievement).filter(
        _proj_or_filter(models.Achievement, project_id, proj_name)
    )
    achs = ach_q.order_by(models.Achievement.updated_at.desc()).all()

    # ── UpdateSubmissions ──────────────────────────────────
    sub_q = db.query(models.UpdateSubmission).filter(
        models.UpdateSubmission.project_id == project_id
    ).order_by(models.UpdateSubmission.created_at.desc())

    if is_member:
        sub_q = sub_q.filter(models.UpdateSubmission.submitter == context["name"])
    elif is_coord and not is_owner and not is_super:
        # coordinator 只看转交给自己的待统筹反馈事项
        sub_q = sub_q.filter(
            models.UpdateSubmission.confirm_status.in_(list(SS.WAITING_COORDINATOR_FEEDBACK))
        )

    subs = sub_q.all()

    # 全量提交（用于所有人可见的统计，如 CEO 决策计数）
    all_subs = (
        db.query(models.UpdateSubmission)
        .filter(models.UpdateSubmission.project_id == project_id)
        .all()
    )

    # ── Meetings ───────────────────────────────────────────
    meeting_q = db.query(models.Meeting).filter(
        _proj_or_filter(models.Meeting, project_id, proj_name, str_col="related_special_project")
    )
    meeting_count = meeting_q.count()

    # ── 角色决定各类数据可见性 ─────────────────────────────
    # super_admin / owner: 完整数据
    # project_ceo: 可看整体进展、风险、成果、CEO决策；不看成员提交明细
    # coordinator: 只看转交给自己的待统筹反馈事项（sub_q 已过滤）
    # member: 只看自己任务和提交

    can_see_submissions  = is_super or is_owner or is_global_ceo
    can_see_decisions    = is_super or is_owner or is_project_ceo or is_global_ceo
    can_see_risks        = is_super or is_owner or is_coord or is_project_ceo or is_global_ceo

    # 本月重点：plan_time 包含当前月、且未完成/未暂缓的任务
    # 超期任务排在前面，其次进行中，最后未开始
    _non_terminal = _NOT_STARTED | _IN_PROGRESS | _DELAYED
    def _month_task_priority(t):
        s = t.status or ""
        overdue = s in _active and _is_overdue(t.plan_time or "")
        if s in _DELAYED or overdue: return 0
        if s in _IN_PROGRESS:        return 1
        return 2
    _active = _NOT_STARTED | _IN_PROGRESS
    month_tasks = sorted(
        [t for t in tasks if _plan_time_in_current_month(t.plan_time or "") and (t.status or "") in _non_terminal],
        key=_month_task_priority,
    )
    # 本月无数据时回退到最近更新的任务
    _source_tasks = month_tasks if month_tasks else tasks[:5]

    def _task_dict_with_overdue(t):
        d = crud.to_dict(t)
        d["is_overdue"] = (t.status or "") in _active and _is_overdue(t.plan_time or "")
        return d

    recent_tasks  = [_task_dict_with_overdue(t) for t in _source_tasks[:5]]

    # 独立的延期任务列表：状态为"延期"或已超期（不受本月过滤限制）
    delayed_tasks = [
        _task_dict_with_overdue(t) for t in tasks
        if (t.status or "") in _DELAYED or (
            (t.status or "") in _active and _is_overdue(t.plan_time or "")
        )
    ][:10]
    recent_issues = [crud.to_dict(i) for i in issues[:5]]
    recent_subs   = [crud.to_dict(s) for s in subs[:10]] if can_see_submissions else []
    recent_achs   = [crud.to_dict(a) for a in achs[:10]]

    decisions_list = [crud.to_dict(i) for i in ceo_issues[:20]] if can_see_decisions else []
    risks_list     = [crud.to_dict(i) for i in open_issues[:20]] if can_see_risks else []

    task_s  = _task_stats(tasks)
    sub_s   = _submission_stats(all_subs if can_see_submissions else subs)
    ceo_s   = _ceo_decision_stats(all_subs)

    # coordinator 裁剪——只展示自身相关统计
    if is_coord and not is_owner and not is_super:
        sub_s = {
            "total_submissions":           sub_s["total_submissions"],
            "pending_owner_confirmation":   None,   # coordinator 不看 owner 待确认队列
            "returned_submissions":         None,   # coordinator 不看打回明细
            "confirmed_submissions":        sub_s["confirmed_submissions"],
            "waiting_coordinator":          len(subs),   # sub_q 已收窄到 WAITING_COORDINATOR_FEEDBACK
        }

    # 5D：project_ceo 裁剪——不显示 owner 待确认队列和打回明细
    if is_project_ceo and not is_owner and not is_super:
        sub_s = {
            "total_submissions":          sub_s["total_submissions"],
            "pending_owner_confirmation":  None,   # CEO 不看 owner 待确认队列
            "returned_submissions":        None,   # CEO 不看打回明细
            "confirmed_submissions":       sub_s["confirmed_submissions"],
        }

    completion_rate = round(task_s["completed"] / task_s["total_tasks"] * 100) if task_s["total_tasks"] else 0

    # ── 角色队列：仪表盘右下面板按角色展示不同内容 ────────────
    if is_project_ceo and not is_owner and not is_super:
        queue_type  = "pending_decisions"
        queue_items = [s for s in all_subs if (s.confirm_status or "") in SS.WAITING_CEO_DECISION]
    elif is_global_ceo and not is_owner and not is_super:
        queue_type  = "pending_decisions"
        queue_items = [s for s in all_subs if (s.confirm_status or "") in SS.WAITING_CEO_DECISION]
    elif is_owner or is_super:
        queue_type  = "pending_review"
        queue_items = [s for s in all_subs if (s.confirm_status or "") in SS.PENDING_OWNER_REVIEW]
    elif is_coord:
        queue_type  = "pending_coordinator"
        queue_items = [s for s in all_subs if (s.confirm_status or "") in SS.WAITING_COORDINATOR_FEEDBACK]
    else:
        queue_type  = "in_progress"
        queue_items = [s for s in subs if (s.confirm_status or "") not in SS.ALL_TERMINAL]

    role_queue = {
        "type":  queue_type,
        "count": len(queue_items),
        "items": [crud.to_dict(s) for s in queue_items[:10]],
    }

    return {
        "project": {"id": project_id, "name": proj_name or ""},
        "access": {
            "can_view_decisions":           can_see_decisions,
            "can_view_risks":               can_see_risks,
            "can_view_confirmation_center": can_see_submissions,
            "can_view_settings":            bool(context.get("can_view_settings")),
            "effective_roles":              sorted(roles - {"super_admin"}),
        },
        # ── 新增详细统计字段 ──────────────────────────────
        "task_stats": task_s,
        "achievement_stats": {
            "total_achievements": len(achs),
            "recent_achievements": recent_achs[:5],
        },
        "issue_stats": {
            "total_issues":          len(issues),
            "open_issues":           len(open_issues),
            "high_priority_issues":  len(high_pri),
            "waiting_ceo_decision":  len(ceo_issues) if can_see_decisions else 0,
        },
        "submission_stats":   sub_s,
        "ceo_decision_stats": ceo_s,
        "meeting_count":      meeting_count,
        "recent": {
            "submissions":  recent_subs,
            "tasks":        recent_tasks,
            "issues":       recent_issues,
            "delayed_tasks": delayed_tasks,
        },
        # ── Legacy fields（保留旧前端兼容性）───────────────
        "summary": {
            "task_count":                 task_s["total_tasks"],
            "achievement_count":          len(achs),
            "open_issue_count":           len(open_issues),
            "pending_confirmation_count": sub_s.get("pending_owner_confirmation") or 0,
        },
        "project_cards": [{
            "special_project":  proj_name or "",
            "project_id":       project_id,
            "owners":           "、".join(sorted({t.owner for t in tasks if t.owner})),
            "task_count":       task_s["total_tasks"],
            "completed_count":  task_s["completed"],
            "completion_rate":  completion_rate,
            "achievement_count": len(achs),
            "open_issue_count": len(open_issues),
        }],
        "status_stats": {
            s: task_s.get(s, 0)
            for s in ("not_started", "in_progress", "completed", "delayed", "paused")
        },
        "pending_tasks":      [crud.to_dict(t) for t in tasks if (t.status or "") not in _COMPLETED][:10],
        "decisions":          decisions_list,
        "risks":              risks_list,
        "latest_achievements": recent_achs,
        "role_queue":         role_queue,
    }


def _global_role_queue(context: dict, db: Session) -> dict:
    """全局模式下按用户全局角色返回角色队列。"""
    is_super = context.get("can_view_all") or context.get("is_tech_admin")
    is_ceo   = context.get("is_ceo")
    owned    = bool(context.get("owned_projects"))
    coord    = bool(context.get("coordinated_projects"))

    if is_super or is_ceo:
        queue_type = "pending_decisions"
        q = db.query(models.UpdateSubmission).filter(
            models.UpdateSubmission.confirm_status.in_(SS.WAITING_CEO_DECISION)
        )
    elif owned:
        queue_type = "pending_review"
        owned_names = context.get("owned_projects", [])
        owned_proj_ids: list[int] = []
        if owned_names:
            rows = db.query(models.Project.id).filter(models.Project.name.in_(owned_names)).all()
            owned_proj_ids = [r[0] for r in rows]
        q = db.query(models.UpdateSubmission).filter(
            models.UpdateSubmission.confirm_status.in_(SS.PENDING_OWNER_REVIEW),
        )
        if owned_proj_ids:
            q = q.filter(models.UpdateSubmission.project_id.in_(owned_proj_ids))
    elif coord:
        queue_type = "pending_coordinator"
        coord_names = context.get("coordinated_projects", [])
        coord_proj_ids: list[int] = []
        if coord_names:
            rows = db.query(models.Project.id).filter(models.Project.name.in_(coord_names)).all()
            coord_proj_ids = [r[0] for r in rows]
        q = db.query(models.UpdateSubmission).filter(
            models.UpdateSubmission.confirm_status.in_(SS.WAITING_COORDINATOR_FEEDBACK),
        )
        if coord_proj_ids:
            q = q.filter(models.UpdateSubmission.project_id.in_(coord_proj_ids))
    else:
        queue_type = "in_progress"
        q = db.query(models.UpdateSubmission).filter(
            models.UpdateSubmission.submitter == context["name"],
            ~models.UpdateSubmission.confirm_status.in_(SS.ALL_TERMINAL),
        )

    items = q.order_by(models.UpdateSubmission.created_at.desc()).limit(10).all()
    return {
        "type":  queue_type,
        "count": q.count(),
        "items": [crud.to_dict(s) for s in items],
    }


# ── 全局模式辅助 ──────────────────────────────────────────────

def _global_recent_tasks(tasks: list) -> list:
    """
    全局模式下的"本月重点"：
    与 _project_overview 保持一致——取 plan_time 包含当前月且未完成/未暂缓的任务，
    超期/延期优先，回退到最近更新的前 5 条。
    """
    _active      = _NOT_STARTED | _IN_PROGRESS
    _non_terminal = _NOT_STARTED | _IN_PROGRESS | _DELAYED

    def _priority(t):
        s = t.status or ""
        overdue = s in _active and _is_overdue(t.plan_time or "")
        if s in _DELAYED or overdue: return 0
        if s in _IN_PROGRESS:        return 1
        return 2

    def _with_overdue(t):
        d = crud.to_dict(t)
        d["is_overdue"] = (t.status or "") in _active and _is_overdue(t.plan_time or "")
        return d

    month_tasks = sorted(
        [t for t in tasks if _plan_time_in_current_month(t.plan_time or "") and (t.status or "") in _non_terminal],
        key=_priority,
    )
    source = month_tasks if month_tasks else tasks[:5]
    return [_with_overdue(t) for t in source[:5]]


def _global_delayed_tasks(tasks: list) -> list:
    """全局模式下的延期任务列表：状态为"延期"或已超期（不受本月过滤限制）。"""
    _active = _NOT_STARTED | _IN_PROGRESS

    def _with_overdue(t):
        d = crud.to_dict(t)
        d["is_overdue"] = (t.status or "") in _active and _is_overdue(t.plan_time or "")
        return d

    return [
        _with_overdue(t) for t in tasks
        if (t.status or "") in _DELAYED or (
            (t.status or "") in _active and _is_overdue(t.plan_time or "")
        )
    ][:10]


# ── 全局模式 overview（原逻辑不变）──────────────────────────

def _global_overview(
    context: dict,
    special_project_filter: str | None,
    owner_filter: str | None,
    status_filter: str | None,
    month_filter: str | None,
    db: Session,
) -> dict:
    # 合并新旧两套权限系统，得到当前用户可见的 project_id 列表
    visible_proj_ids = _get_visible_project_ids(context, db)

    task_q = _apply_project_scope(db.query(models.Task), context, models.Task, visible_proj_ids).filter(models.Task.is_deleted == False)
    # 只统计活跃（未归档）项目的任务
    if not special_project_filter:
        active_proj_names = db.query(models.Project.name).filter(models.Project.status != "archived")
        task_q = task_q.filter(models.Task.special_project.in_(active_proj_names))
    if special_project_filter:
        task_q = task_q.filter(models.Task.special_project == special_project_filter)
    if owner_filter:
        task_q = task_q.filter(models.Task.owner == owner_filter)
    if status_filter:
        task_q = task_q.filter(models.Task.status == status_filter)
    if month_filter:
        task_q = task_q.filter(models.Task.plan_time.like(f"{_month_to_iso(month_filter)}%"))
    tasks = task_q.all()

    # ── project_cards：从实际项目列表构建，兼容新旧数据 ──────────
    # 不再依赖任务的 special_project 分组（新系统任务 special_project 可能为空）
    if visible_proj_ids is None:
        visible_projs = (
            db.query(models.Project)
            .filter(models.Project.status != "archived")
            .order_by(models.Project.sort_order, models.Project.id)
            .all()
        )
    elif visible_proj_ids:
        visible_projs = (
            db.query(models.Project)
            .filter(models.Project.id.in_(visible_proj_ids), models.Project.status != "archived")
            .order_by(models.Project.sort_order, models.Project.id)
            .all()
        )
    else:
        visible_projs = []

    project_cards = []
    for proj in visible_projs:
        p_task_q = db.query(models.Task).filter(_proj_or_filter(models.Task, proj.id, proj.name), models.Task.is_deleted == False)
        if month_filter:
            p_task_q = p_task_q.filter(models.Task.plan_time.like(f"{month_filter}%"))
        p_tasks = p_task_q.all()
        completed = sum(1 for t in p_tasks if (t.status or "") in _COMPLETED)
        total = len(p_tasks)
        ach_count = (
            db.query(models.Achievement)
            .filter(_proj_or_filter(models.Achievement, proj.id, proj.name))
            .count()
        )
        open_issue_count = (
            db.query(models.Issue)
            .filter(
                _proj_or_filter(models.Issue, proj.id, proj.name),
                models.Issue.status.in_(["待处理", "处理中"]),
            )
            .count()
        )
        latest = max((t.updated_at for t in p_tasks if t.updated_at), default=None)
        project_cards.append({
            "special_project":   proj.name,
            "name":              proj.name,
            "project_id":        proj.id,
            "owners":            "、".join(sorted({t.owner for t in p_tasks if t.owner})),
            "task_count":        total,
            "completed_count":   completed,
            "completion_rate":   round(completed / total * 100) if total else 0,
            "achievement_count": ach_count,
            "open_issue_count":  open_issue_count,
            "latest_update":     latest.isoformat(timespec="seconds") if latest else "",
        })

    status_stats = Counter(t.status for t in tasks)
    pending_tasks = [crud.to_dict(t) for t in tasks if t.status != "已完成"][:10]

    issue_q    = _apply_project_scope(db.query(models.Issue), context, models.Issue, visible_proj_ids).order_by(models.Issue.updated_at.desc())
    issue_rows = issue_q.all()
    decisions, risks = [], []
    for row in issue_rows:
        is_decision = bool(row.need_decision_by) or "决策" in (row.issue_type or "")
        if is_decision and can_view_issue_decisions(context):
            decisions.append(crud.to_dict(row))
        elif not is_decision and can_view_issue_risks(context):
            risks.append(crud.to_dict(row))

    latest_achievements_q = (
        _apply_project_scope(db.query(models.Achievement), context, models.Achievement, visible_proj_ids)
        .order_by(models.Achievement.updated_at.desc())
    )
    latest_achievements = [crud.to_dict(a) for a in latest_achievements_q.limit(10).all()]

    pending_confirmation_count = 0
    if can_access_confirmation_center(context):
        pending_confirmation_count = (
            db.query(models.UpdateSubmission)
            .filter(models.UpdateSubmission.confirm_status.in_(SS.ALL_ACTIVE))
            .count()
        )

    visible_achievement_count = _apply_project_scope(
        db.query(models.Achievement), context, models.Achievement, visible_proj_ids
    ).count()
    visible_open_issue_count = (
        _apply_project_scope(db.query(models.Issue), context, models.Issue, visible_proj_ids)
        .filter(models.Issue.status.in_(["待处理", "处理中"]))
        .count()
    )

    task_s = _task_stats(tasks)

    return {
        "access": {
            "can_view_decisions":           can_view_issue_decisions(context),
            "can_view_risks":               can_view_issue_risks(context),
            "can_view_confirmation_center": can_access_confirmation_center(context),
            "can_view_settings":            context.get("can_view_settings", False),
        },
        "filters": {
            "projects":  [p["special_project"] for p in project_cards],
            "owners":    [p[0] for p in db.query(models.Task.owner).distinct().all() if p[0]],
            "statuses":  ["未开始", "推进中", "已完成", "延期", "暂缓"],
        },
        "summary": {
            "task_count":                 len(tasks),
            "achievement_count":          visible_achievement_count,
            "open_issue_count":           visible_open_issue_count,
            "pending_confirmation_count": pending_confirmation_count,
        },
        # ── 与项目模式对齐的字段（供新前端仪表盘使用）──────────
        "task_stats": task_s,
        "achievement_stats": {
            "total_achievements": visible_achievement_count,
            "recent_achievements": latest_achievements[:5],
        },
        "issue_stats": {
            "total_issues":         len(issue_rows),
            "open_issues":          visible_open_issue_count,
            "high_priority_issues": sum(1 for i in issue_rows if (i.priority or "") == "高" and (i.status or "") in ("待处理", "处理中")),
            "waiting_ceo_decision": sum(1 for i in issue_rows if bool(i.need_decision_by) or "决策" in (i.issue_type or "")) if can_view_issue_decisions(context) else 0,
        },
        "recent": {
            "submissions":   [],
            "tasks":         _global_recent_tasks(tasks),
            "issues":        [crud.to_dict(i) for i in issue_rows[:5]],
            "delayed_tasks": _global_delayed_tasks(tasks),
        },
        "project_cards":     project_cards,
        "status_stats":      dict(status_stats),
        "pending_tasks":     pending_tasks,
        "decisions":         decisions if can_view_issue_decisions(context) else [],
        "risks":             risks if can_view_issue_risks(context) else [],
        "latest_achievements": latest_achievements,
        "role_queue":        _global_role_queue(context, db),
    }


# ── 端点 ─────────────────────────────────────────────────────

@router.get("/overview")
def overview(
    project_id: int | None = None,
    special_project: str | None = None,
    owner: str | None = None,
    status: str | None = None,
    month: str | None = None,
    current_user: str = Depends(get_current_user_name),
    db: Session = Depends(get_db),
):
    current_user = require_login(current_user, db)
    context = get_user_context_from_db(current_user, db)

    # ── 解析 effective_project_id ────────────────────────
    resolution = resolve_project_context(
        db,
        project_id=project_id,
        special_project=special_project,
    )
    effective_project_id: int | None = resolution["project_id"]
    proj_name_label: str = resolution["project_name"] or (special_project or "")

    if project_id is not None and not resolution["is_valid"]:
        raise HTTPException(404, "project not found")
    if project_id is None and special_project and effective_project_id is None:
        # 无法解析：返回空统计，不报 500，避免旧前端崩溃
        _require_global_read_scope(context)
        return _empty_project_overview(context, special_project)

    # ── 项目模式 ──────────────────────────────────────────
    if effective_project_id is not None:
        require_project_access(current_user, effective_project_id, db)
        # 查项目名（用于旧数据 OR 过滤和 project_cards 展示）
        proj_name = resolve_project_context(db, project_id=effective_project_id)["project_name"] or proj_name_label
        return _project_overview(context, effective_project_id, proj_name, owner, status, month, db)

    # ── 全局模式（无 project_id / special_project）────────
    # 完全保留原有行为，保证旧前端不受影响
    _require_global_read_scope(context)
    return _global_overview(context, None, owner, status, month, db)


# ── 导出周报 ──────────────────────────────────────────────────

def _build_weekly_report(
    context: dict,
    project_id: int | None,
    month: str | None,
    db: Session,
) -> bytes:
    """生成 Word 格式周报，返回 bytes。"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx 未安装")

    today = date.today()
    report_month = month or f"{today.year}年{today.month}月"

    # ── 拉数据 ───────────────────────────────────────────────
    visible_proj_ids = _get_visible_project_ids(context, db)
    task_q = _apply_project_scope(db.query(models.Task), context, models.Task, visible_proj_ids).filter(models.Task.is_deleted == False)
    if project_id is not None:
        proj_name = resolve_project_context(db, project_id=project_id)["project_name"] or ""
        task_q = task_q.filter(_proj_or_filter(models.Task, project_id, proj_name))
    if month:
        task_q = task_q.filter(models.Task.plan_time.like(f"{_month_to_iso(month)}%"))
    tasks = task_q.all()

    ach_q = _apply_project_scope(db.query(models.Achievement), context, models.Achievement, visible_proj_ids)
    if project_id is not None:
        ach_q = ach_q.filter(_proj_or_filter(models.Achievement, project_id, proj_name))
    achs = ach_q.order_by(models.Achievement.updated_at.desc()).limit(20).all()

    issue_q = _apply_project_scope(db.query(models.Issue), context, models.Issue, visible_proj_ids)
    if project_id is not None:
        issue_q = issue_q.filter(_proj_or_filter(models.Issue, project_id, proj_name))
    issues = issue_q.filter(models.Issue.status.in_(["待处理", "处理中"])).order_by(models.Issue.updated_at.desc()).limit(20).all()

    stats = _task_stats(tasks)

    # ── 按专项分组 ──────────────────────────────────────────
    grouped: dict[str, list] = defaultdict(list)
    for t in tasks:
        grouped[t.special_project or "（未分类）"].append(t)

    # ── 构建 Word 文档 ──────────────────────────────────────
    doc = Document()

    # 标题
    title = doc.add_heading(f"博维咨询 · AI项目进展周报", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"报告月份：{report_month}　　生成日期：{today.strftime('%Y年%m月%d日')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 一、整体概况
    doc.add_heading("一、整体任务概况", level=1)
    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Light Shading Accent 1"
    headers = ["总任务数", "未启动", "进行中", "已完成", "延期/超期"]
    values  = [stats["total_tasks"], stats["not_started"], stats["in_progress"],
               stats["completed"], stats["delayed"]]
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        tbl.rows[1].cells[i].text = str(values[i])
    doc.add_paragraph()

    # 二、各专项进度
    doc.add_heading("二、各专项任务进度", level=1)
    if grouped:
        for proj, rows in sorted(grouped.items()):
            done  = sum(1 for t in rows if (t.status or "") in _COMPLETED)
            total = len(rows)
            rate  = round(done / total * 100) if total else 0
            doc.add_heading(f"▶ {proj}（{done}/{total}，完成率 {rate}%）", level=2)
            for t in rows:
                status = t.status or "未知"
                owner  = t.owner or "-"
                plan   = t.plan_time or "-"
                overdue_tag = "【超期】" if (status in _NOT_STARTED | _IN_PROGRESS and _is_overdue(t.plan_time or "")) else ""
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{overdue_tag}{t.key_task or '（无标题）'}　[{status}]　负责人：{owner}　计划：{plan}")
                if overdue_tag:
                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    else:
        doc.add_paragraph("（本月暂无任务数据）")
    doc.add_paragraph()

    # 三、近期成果
    doc.add_heading("三、近期成果", level=1)
    if achs:
        for a in achs[:10]:
            doc.add_paragraph(f"✓ {a.name or '（无标题）'}　[{a.achievement_type or '-'}]　负责人：{a.owner or '-'}", style="List Bullet")
    else:
        doc.add_paragraph("（暂无成果记录）")
    doc.add_paragraph()

    # 四、待处理问题与风险
    doc.add_heading("四、待处理问题与风险", level=1)
    if issues:
        for iss in issues[:10]:
            priority = iss.priority or "普通"
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{priority}] {iss.description or '（无描述）'}　专项：{iss.special_project or '-'}")
            if priority == "高":
                run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    else:
        doc.add_paragraph("（暂无待处理问题）")

    # 保存到内存
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.get("/export-weekly-report")
def export_weekly_report(
    project_id: int | None = None,
    month: str | None = None,
    current_user: str = Depends(get_current_user_name),
    db: Session = Depends(get_db),
):
    current_user = require_login(current_user, db)
    context = get_user_context_from_db(current_user, db)
    today = date.today()
    report_month = month or f"{today.year}年{today.month}月"

    # 解析 project_id
    effective_pid: int | None = project_id
    if effective_pid is None:
        _require_global_read_scope(context)
    else:
        require_project_access(current_user, effective_pid, db)
    try:
        data = _build_weekly_report(context, effective_pid, month, db)
    except RuntimeError as e:
        from fastapi import HTTPException
        raise HTTPException(500, str(e))

    from urllib.parse import quote
    safe_month = report_month.replace("年", "-").replace("月", "")
    filename = f"周报_{safe_month}.docx"
    encoded_filename = quote(filename, safe="")
    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )

